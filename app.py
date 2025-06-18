
# Hugging Face healthcheck workaround
import os

if os.environ.get("HF_SPACE_ID"):
    import socket
    import http.server
    import threading

    class HealthHandler(http.server.BaseHTTPRequestHandler):
        def do_GET(self):
            self.send_response(200)
            self.send_header("Content-type", "text/plain")
            self.end_headers()
            self.wfile.write(b"ok")

    def run_health_server():
        server_address = ('', 7861)
        httpd = http.server.HTTPServer(server_address, HealthHandler)
        httpd.serve_forever()

    thread = threading.Thread(target=run_health_server)
    thread.daemon = True
    thread.start()

# --- Actual App Starts Below ---

import streamlit as st
import json
import re
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from io import BytesIO

# --- Load OpenAI credentials from Streamlit secrets ---
api_key = st.secrets["OPENAI_API_KEY"]
project_id = st.secrets["OPENAI_PROJECT_ID"]
MODEL = "gpt-4o"

SECTORS = [
    "Retail & Leisure",
    "Industrials",
    "Technology, Media and Telecoms",
    "Financial & Professional Services",
    "Real Estate & Construction",
    "Energy, Chemicals, Mining & Utilities",
    "Healthcare & Pharma"
]

def get_client():
    return OpenAI(api_key=api_key, project=project_id)

def clean_summary_text(text):
    # Remove markdown-style bold/italics/underscores
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    # Remove stray (Link) from GPT
    text = re.sub(r"\(Link\)", "", text)
    # Fix spacing between numbers and words (e.g. 2.6billion → 2.6 billion)
    text = re.sub(r"(\d)([A-Za-z])", r"\1 \2", text)
    # Flatten newlines and trim whitespace
    text = text.replace("\n", " ").replace("\r", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def format_summary(company, summary_text):
    summary_clean = clean_summary_text(summary_text)
    dash_index = summary_clean.find("–")
    if dash_index == -1:
        dash_index = summary_clean.find("-")
    if dash_index != -1:
        body = summary_clean[dash_index + 1:].strip()
        if body and not body[0].isupper():
            body = body[0].lower() + body[1:]
        return f"**{company}** – {body} (Link)"
    else:
        return f"**{company}** – {summary_clean} (Link)"

def generate_summary(rns_text):
    client = get_client()
    prompt = f"""
You're a financial journalist. Summarise this UK RNS announcement in one bullet point.

Editorial rules:
- Begin with the company name in **bold**, followed by an en dash (–)
- Do not repeat the company name in the body
- Use "has announced" for company-led updates
- Use "has said that" for third-party or external developments
- Begin the sentence after the dash in lowercase, unless it's a proper noun
- Correctly capitalise initials in names (e.g. "J.T. Starzecki")
- Include only strategic, financial, or operational business facts
DO NOT include (Link) at the end.

RNS:
{rns_text}
"""
    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return clean_summary_text(response.choices[0].message.content.strip())

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def docx_export(summaries):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    grouped = {sector: [] for sector in SECTORS}
    for item in summaries:
        grouped[item["sector"]].append(item)

    for sector in SECTORS:
        entries = sorted(grouped[sector], key=lambda x: x["company"])
        if entries:
            p = doc.add_paragraph()
            run = p.add_run(sector)
            run.bold = True
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 0)

            for item in entries:
                para = doc.add_paragraph()
                summary_clean = clean_summary_text(item["summary"]).replace(" (Link)", "")
                dash_index = summary_clean.find("–")
                if dash_index != -1:
                    summary_part = summary_clean[dash_index + 1:].strip()
                else:
                    summary_part = summary_clean
                para.add_run(item["company"]).bold = True
                para.add_run(" – ")
                para.add_run(summary_part + " ")
                add_hyperlink(para, "(Link)", item["link"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def today():
    return datetime.now().strftime("%Y-%m-%d")

# --- Streamlit UI ---
st.set_page_config(page_title="RNS Summariser Tool", layout="wide")
st.title("📈 RNS Summariser Tool (Formatted Output)")
st.empty()
st.markdown("<!-- Hugging Face healthcheck passthrough -->")

if "summaries" not in st.session_state:
    st.session_state.summaries = []

with st.form("rns_form"):
    rns_text = st.text_area("Paste RNS Text", height=200)
    company = st.text_input("Company Name")
    link = st.text_input("RNS Link (URL)")
    sector = st.selectbox("Sector", SECTORS)
    submitted = st.form_submit_button("Summarise & Add")

    if submitted:
        if rns_text.strip() and company and link:
            try:
                raw_summary = generate_summary(rns_text)
                st.session_state.summaries.append({
                    "company": company,
                    "link": link,
                    "sector": sector,
                    "summary": raw_summary
                })
                st.success(f"✅ Added summary for {company}")
            except Exception as e:
                st.error(f"❌ Error: {e}")
        else:
            st.warning("Please fill in all fields.")

if st.session_state.summaries:
    st.subheader("Summarised Entries")
    grouped = {sector: [] for sector in SECTORS}
    for item in st.session_state.summaries:
        grouped[item["sector"]].append(item)

    for sector in SECTORS:
        entries = sorted(grouped[sector], key=lambda x: x["company"])
        if entries:
            st.markdown(f"### {sector}")
            for item in entries:
                formatted = format_summary(item["company"], item["summary"])
                st.markdown(formatted)
                st.markdown("---")

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "⬇️ Download as JSON",
            data=json.dumps(st.session_state.summaries, indent=2),
            file_name=f"rns_summaries_{today()}.json",
            mime="application/json"
        )
    with col2:
        st.download_button(
            "⬇️ Download as Word (.docx)",
            data=docx_export(st.session_state.summaries),
            file_name=f"rns_summaries_{today()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
