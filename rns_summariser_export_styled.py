import streamlit as st
from openai import OpenAI
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import json

# --- API Setup ---
load_dotenv()
api_key = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
if not api_key:
    st.error("No OpenAI API key found. Add it to .streamlit/secrets.toml or .env.")
    st.stop()

client = OpenAI(api_key=api_key)

# --- Constants ---
SECTIONS = [
    "Retail & Leisure",
    "Industrials",
    "Technology, Media and Telecoms",
    "Financial & Professional Services",
    "Real Estate & Construction",
    "Energy, Chemicals, Mining & Utilities",
    "Healthcare & Pharma"
]

# --- GPT Summary Generator ---
def generate_summary(rns_text):
    prompt = f"""
Write a one-line UK stock market RNS summary in no more than 100 words.

- Start with the company name followed by an en dash (–)
- DO NOT repeat the company name in the body
- Begin the sentence after the dash with a lowercase letter
- Include key facts only: results, revenue, profit/loss, dividends, acquisitions, guidance
- Do NOT include links, URLs or disclaimers
- End with: (Link)

RNS:
{rns_text}
"""
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4
    )
    return response.choices[0].message.content.strip()

# --- Hyperlink Formatter ---
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

# --- Word Export Function ---
def docx_export(summaries):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    grouped = {section: [] for section in SECTIONS}
    for item in summaries:
        grouped[item["sector"]].append(item)

    for section in SECTIONS:
        entries = sorted(grouped[section], key=lambda x: x["company"])
        if entries:
            p = doc.add_paragraph()
            run = p.add_run(section)
            run.bold = True
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 0)

            for item in entries:
                para = doc.add_paragraph()
                cleaned = item["summary"].replace("(Link)", "").strip()
                dash_index = cleaned.find("–")
                if dash_index != -1 and dash_index + 2 < len(cleaned):
                    cleaned = cleaned[:dash_index+2] + cleaned[dash_index+2].lower() + cleaned[dash_index+3:]
                para.add_run(cleaned + " ")
                add_hyperlink(para, "(Link)", item["link"])

    path = "rns_summary_output.docx"
    doc.save(path)
    with open(path, "rb") as f:
        return f.read()

def today():
    return datetime.now().strftime("%Y-%m-%d")

# --- Session State ---
if "summaries" not in st.session_state:
    st.session_state.summaries = []

# --- UI ---
st.title("📈 RNS Summariser Tool")

with st.form("rns_form"):
    rns_text = st.text_area("Paste RNS Text", height=200)
    company = st.text_input("Company Name")
    link = st.text_input("RNS Link (URL)")
    sector = st.selectbox("Sector", SECTIONS)
    submitted = st.form_submit_button("Summarise & Add")

    if submitted:
        if rns_text.strip() and company and link:
            summary = generate_summary(rns_text)
            st.session_state.summaries.append({
                "company": company,
                "link": link,
                "sector": sector,
                "summary": summary
            })
            st.success(f"✅ Added summary for {company}")
        else:
            st.warning("Please fill in all fields.")

# --- Display Summaries ---
if st.session_state.summaries:
    st.subheader("Summarised Entries")
    grouped = {section: [] for section in SECTIONS}
    for item in st.session_state.summaries:
        grouped[item["sector"]].append(item)

    for section in SECTIONS:
        entries = sorted(grouped[section], key=lambda x: x["company"])
        if entries:
            st.markdown(f"### {section}")
            for item in entries:
                summary = item["summary"].replace("(Link)", "").strip()
                dash_index = summary.find("–")
                if dash_index != -1 and dash_index + 2 < len(summary):
                    summary = summary[:dash_index+2] + summary[dash_index+2].lower() + summary[dash_index+3:]
                st.markdown(f"{summary} [**(Link)**]({item['link']})")
                st.markdown("---")

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "⬇️ Download as JSON",
            data=json.dumps(st.session_state.summaries, indent=2),
            file_name=f"rns_summaries_{today()}.json",
            mime="application/json"
        )
    with col2:
        st.download_button(
            "⬇️ Download as Word (.docx)",
            data=docx_export(st.session_state.summaries),
            file_name=f"rns_summaries_{today()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
